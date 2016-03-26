import requests, time, operator
from bs4 import *
from selenium import webdriver
from docx import *
#getting list of organisations
res1 = requests.get('https://developers.google.com/open-source/gsoc/2016/organizations')
soup1 = BeautifulSoup(res1.text,'html.parser')
name_tag = soup1.find('div',{'class':'devsite-article-body clearfix'})
list_tag = name_tag.find('ul')

doc = Document()
p = doc.add_paragraph()
p.add_run('Accepted organisations :\n\n').bold = True

links = []
numbering = 1

for name in list_tag.findAll('li'):
    p.add_run(str(numbering)+'. '+name.string+'\n')
    numbering += 1
    link = name.find('a')
    link = str(link)
    end = link.find('>')
    link = link[9:end-2]
    link += '/'
    links.append(link)

doc.add_page_break()

results = dict()

#mention the path of phantomjs.exe in the paranthesis
driver = webdriver.PhantomJS('C:/phantomjs.exe')

p = doc.add_paragraph()
p.add_run("Technologies list (individual companies) :\n\n").bold = True

for link_name in links:
    driver.get(link_name)
    #time delay to load the page properly, otherwise page is returned with null
    time.sleep(2)
    try:
        for link in driver.find_elements_by_tag_name('div'):
            page = link.text
            break
        
        text = page.split('\n')

        for link in driver.find_elements_by_tag_name('h1'):
            name_of_org = link.text
            break

        p.add_run('\t'+str(name_of_org)+'\n').bold = True
        
        for z1 in range(len(text)):
            if(text[z1] == 'Technologies'):
                start = z1
                start += 1
            if(text[z1] == 'Topics'):
                end = z1
                break
            
        for i in range(start,end):
            if(text[i] not in results):
                results.update({text[i]:1})
            else:
                count = results[text[i]]
                count += 1
                results.update({text[i]:count})
            p.add_run('\t> '+str(text[i])+'\n')
        p.add_run('\n')
    except:
        pass

doc.add_page_break()

driver.quit()

#sort the technologies list according to the related number of organisations
results = sorted(results.items(), key = operator.itemgetter(1))
results.reverse()

p = doc.add_paragraph()
p.add_run("Popularity of Technologies : \t[Technology - Related number of organisations]\n").bold = True
t = doc.add_table(len(results),2,"TableGrid")

for i in range(len(results)):
    x = results[i]
    row = t.rows[i]
    row.cells[0].text = str(x[0])
    row.cells[1].text = str(x[1])

doc.save('E:/gsoc_data.docx')
